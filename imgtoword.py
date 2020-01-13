from PIL import Image
from docx import Document

basewidth = 300
img = Image.open('yael.jpg')
wpercent = (basewidth/float(img.size[0]))
hsize = int((float(img.size[1])*float(wpercent)))
img = img.resize((basewidth,hsize), Image.ANTIALIAS)
img.save('yaelgood.jpg')
document = Document()
p = document.add_paragraph()
r = p.add_run()
r.add_picture('yaelgood.jpg')
document.save('withImage.docx')

document.add_picture("yaelgood.jpg", width=(2.0), height=(0.2))





